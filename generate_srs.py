#!/usr/bin/env python3
"""
Script to generate SRS (Software Requirements Specification) Document
for the Rental Management System in DOCX format.

This document follows the IEEE 830-1998 standard for SRS documents.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_toc(doc):
    """Add Table of Contents placeholder"""
    doc.add_paragraph('Table of Contents', style='Heading 1')
    doc.add_paragraph('(This section will be auto-generated in Word: References > Table of Contents)')
    doc.add_paragraph()
    
    # Add manual TOC structure for reference
    toc_items = [
        '1.0 Introduction',
        '    1.1 Purpose',
        '    1.2 Scope of Project',
        '    1.3 Glossary',
        '    1.4 References',
        '    1.5 Overview of Document',
        '    1.6 Stakeholders',
        '2.0 Overall Description',
        '    2.1 System Environment',
        '    2.2 Functional Requirements Specification',
        '    2.3 User Characteristics',
        '    2.4 Non-Functional Requirements',
        '    2.5 Budget',
        '    2.6 Feasibility Study',
        '3.0 Requirements Specification',
        '    3.1 External Interface Requirements',
        '    3.2 Functional Requirements',
        '    3.3 Detailed Non-Functional Requirements',
        '    3.4 Data Model',
        'Appendix A: PlantUML Diagram Instructions',
    ]
    
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()

def add_figures_list(doc):
    """Add List of Figures"""
    doc.add_paragraph('List of Figures', style='Heading 1')
    figures = [
        'Figure 1 - System Environment (Use Case Diagram)',
        'Figure 2 - User Roles and System Architecture',
        'Figure 3 - Admin Use Cases',
        'Figure 4 - Landlord Use Cases',
        'Figure 5 - Tenant Use Cases',
        'Figure 6 - Data Model Structure (Class Diagram)',
        'Figure 7 - Lease Management Process (State Diagram)',
        'Figure 8 - Payment Processing Flow (Activity Diagram)'
    ]
    for fig in figures:
        doc.add_paragraph(fig, style='List Bullet')
    doc.add_page_break()

def add_introduction(doc):
    """Add Section 1.0 - Introduction"""
    doc.add_heading('1.0 Introduction', level=1)
    
    # 1.1 Purpose
    doc.add_heading('1.1 Purpose', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'This Software Requirements Specification (SRS) document provides a complete '
        'description of the Rental Management System. It describes the functional and '
        'non-functional requirements for the system, which will be used by property managers, '
        'landlords, and tenants to manage rental properties, leases, payments, and related '
        'activities.\n\n'
        'This document is intended for:\n'
    )
    doc.add_paragraph('System developers and engineers', style='List Bullet')
    doc.add_paragraph('Project managers and stakeholders', style='List Bullet')
    doc.add_paragraph('Quality assurance and testing teams', style='List Bullet')
    doc.add_paragraph('Future maintenance teams', style='List Bullet')
    
    # 1.2 Scope
    doc.add_heading('1.2 Scope of Project', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The Rental Management System is a full-stack web application designed to streamline '
        'property rental operations. The system will provide:\n'
    )
    doc.add_paragraph('Property and unit management capabilities', style='List Bullet')
    doc.add_paragraph('Tenant profile and lease agreement management', style='List Bullet')
    doc.add_paragraph('Payment tracking and financial reporting', style='List Bullet')
    doc.add_paragraph('Document storage and management', style='List Bullet')
    doc.add_paragraph('Communication system between landlords and tenants', style='List Bullet')
    doc.add_paragraph('Double-entry accounting module', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run(
        '\nThe system supports three user roles:\n'
    )
    doc.add_paragraph('Admin - Full system access and user management', style='List Bullet')
    doc.add_paragraph('Landlord - Property and tenant management', style='List Bullet')
    doc.add_paragraph('Tenant - Lease viewing and payment submission', style='List Bullet')
    
    # 1.3 Glossary
    doc.add_heading('1.3 Glossary', level=2)
    glossary_terms = [
        ('Admin', 'System administrator with full access to all features'),
        ('API', 'Application Programming Interface for frontend-backend communication'),
        ('JWT', 'JSON Web Token used for authentication'),
        ('Landlord', 'Property owner who manages rental units'),
        ('Lease', 'Legal agreement between landlord and tenant'),
        ('Property', 'A rental building or complex containing one or more units'),
        ('Tenant', 'Individual renting a unit from a landlord'),
        ('Unit', 'Individual rental space within a property'),
        ('REST', 'Representational State Transfer - API architecture'),
        ('SRS', 'Software Requirements Specification')
    ]
    for term, definition in glossary_terms:
        p = doc.add_paragraph()
        p.add_run(f'{term}: ').bold = True
        p.add_run(definition)
    
    # 1.4 References
    doc.add_heading('1.4 References', level=2)
    references = [
        'IEEE Std 830-1998, IEEE Recommended Practice for Software Requirements Specifications',
        'Django 5.x Documentation - https://docs.djangoproject.com/',
        'Django REST Framework Documentation - https://www.django-rest-framework.org/',
        'React 18 Documentation - https://react.dev/',
        'MySQL 8.0 Reference Manual',
        'JWT Authentication Specification - RFC 7519'
    ]
    for ref in references:
        doc.add_paragraph(ref, style='List Bullet')
    
    # 1.5 Overview
    doc.add_heading('1.5 Overview of Document', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The remainder of this document contains two sections. Section 2 provides an overview '
        'of the system, including the system environment, functional requirements with use cases, '
        'user characteristics, and non-functional requirements. Section 3 provides detailed '
        'specifications for each requirement, including external interfaces, functional requirements, '
        'and detailed non-functional requirements.'
    )
    
    # 1.6 Stakeholders
    doc.add_heading('1.6 Stakeholders', level=2)
    
    doc.add_heading('Primary Stakeholders', level=3)
    p = doc.add_paragraph()
    p.add_run(
        'These are individuals directly supporting, managing, and enabling the platform operations:\n'
    )
    
    stakeholders_primary = [
        ('Eliwahid Mdee', 'Project Lead & Developer', 
         'System architecture and development\nProject management and coordination\nLeveraging GitHub Student Developer Pack benefits\nTechnical implementation and maintenance'),
        ('Development Team', 'Software Engineers & Developers',
         'Backend development (Django)\nFrontend development (React)\nDatabase design and optimization\nQuality assurance and testing\nSecurity implementation'),
    ]
    
    for name, role, responsibilities in stakeholders_primary:
        p = doc.add_paragraph()
        p.add_run(f'{name} – {role}\n').bold = True
        p.add_run(responsibilities)
        doc.add_paragraph()
    
    doc.add_heading('Secondary Stakeholders', level=3)
    p = doc.add_paragraph()
    p.add_run(
        'These are partners and users who support platform functionality and business continuity:\n'
    )
    
    stakeholders_secondary = [
        ('Property Owners & Landlords', 
         'End users who manage rental properties\nProvide real-world requirements and feedback\nTest and validate system functionality'),
        ('Tenants', 
         'End users who rent properties\nProvide user experience feedback\nValidate payment and communication features'),
        ('System Administrators',
         'Manage system operations\nMonitor system health and security\nHandle user support requests'),
    ]
    
    for role, description in stakeholders_secondary:
        p = doc.add_paragraph()
        p.add_run(f'{role}\n').bold = True
        p.add_run(description)
        doc.add_paragraph()
    
    doc.add_page_break()

def add_overall_description(doc):
    """Add Section 2.0 - Overall Description"""
    doc.add_heading('2.0 Overall Description', level=1)
    
    # 2.1 System Environment
    doc.add_heading('2.1 System Environment', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The Rental Management System is a web-based application with a client-server architecture. '
        'The system consists of:\n'
    )
    doc.add_paragraph('React-based frontend (client) running in web browsers', style='List Bullet')
    doc.add_paragraph('Django REST API backend (server) processing business logic', style='List Bullet')
    doc.add_paragraph('MySQL database for data persistence', style='List Bullet')
    doc.add_paragraph('File storage system for documents and images', style='List Bullet')
    
    # Add PlantUML code for Figure 1
    doc.add_heading('Figure 1 - System Environment (Use Case Diagram)', level=3)
    doc.add_paragraph('PlantUML Code:')
    uml_code = '''@startuml
left to right direction

actor "Admin" as Admin
actor "Landlord" as Landlord  
actor "Tenant" as Tenant

package "Rental Management System" {
    rectangle "React Frontend" as Frontend
    rectangle "Django REST API" as Backend
    database "MySQL Database" as DB
    folder "File Storage" as Storage
}

Admin --> Frontend : manages system
Landlord --> Frontend : manages properties
Tenant --> Frontend : views/pays

Frontend --> Backend : API calls (JWT)
Backend --> DB : CRUD operations
Backend --> Storage : file operations

note right of Frontend
  Responsive web interface
  React 18+ with Tailwind CSS
end note

note right of Backend
  Django REST Framework
  JWT Authentication
  Business Logic
end note

@enduml'''
    doc.add_paragraph(uml_code, style='Intense Quote')
    
    # 2.2 Functional Requirements
    doc.add_heading('2.2 Functional Requirements Specification', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'This section outlines the use cases for each user role. The system has three primary '
        'actors: Admin, Landlord, and Tenant. Each actor has specific permissions and capabilities '
        'within the system.'
    )
    
    add_admin_use_cases(doc)
    add_landlord_use_cases(doc)
    add_tenant_use_cases(doc)
    
    # 2.3 User Characteristics
    doc.add_heading('2.3 User Characteristics', level=2)
    
    doc.add_heading('Admin User', level=3)
    p = doc.add_paragraph()
    p.add_run(
        'System administrators have technical knowledge and full system access. They manage '
        'users, monitor system health, and configure system-wide settings.'
    )
    
    doc.add_heading('Landlord User', level=3)
    p = doc.add_paragraph()
    p.add_run(
        'Property owners or managers who need to manage rental properties, tenants, and finances. '
        'They have moderate technical proficiency and business management experience.'
    )
    
    doc.add_heading('Tenant User', level=3)
    p = doc.add_paragraph()
    p.add_run(
        'Individuals renting properties who need basic system access to view lease details, '
        'make payments, and communicate with landlords. Basic computer literacy assumed.'
    )
    
    # 2.4 Non-Functional Requirements
    doc.add_heading('2.4 Non-Functional Requirements', level=2)
    
    nfr_categories = [
        ('Performance', 'System should respond to user actions within 2 seconds under normal load'),
        ('Security', 'JWT-based authentication, role-based access control, encrypted data transmission'),
        ('Usability', 'Intuitive interface, responsive design for mobile and desktop devices'),
        ('Reliability', 'System uptime of 99.5%, automated backups, error recovery mechanisms'),
        ('Scalability', 'Support for up to 10,000 concurrent users and 100,000+ properties'),
        ('Maintainability', 'Modular code structure, comprehensive documentation, automated testing')
    ]
    
    for category, description in nfr_categories:
        p = doc.add_paragraph()
        p.add_run(f'{category}: ').bold = True
        p.add_run(description)
    
    # 2.5 Budget
    doc.add_heading('2.5 Budget (Estimated)', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        'The following budget estimate considers GitHub Student Developer Pack benefits available to '
        'the project lead (Eliwahid Mdee), which provides free access to various development tools '
        'and services:\n'
    )
    
    doc.add_heading('Budget Breakdown', level=3)
    
    # Create a budget table
    budget_items = [
        ('Item', 'Cost (USD)', 'Notes'),
        ('System Development (Backend + Frontend)', '$0', 'In-house development by team'),
        ('UI/UX Design', '$0', 'Using free tools: Figma (free tier), Tailwind CSS'),
        ('Domain Name (.com/.net)', '$12/year', 'Standard domain registration'),
        ('Hosting & Deployment', '$0', 'GitHub Student Pack: Free hosting via Azure, Heroku, or DigitalOcean credits'),
        ('SSL Certificate', '$0', 'Free via Let\'s Encrypt or hosting provider'),
        ('Database Hosting', '$0', 'GitHub Student Pack: Free MySQL hosting via PlanetScale or Railway'),
        ('Development Tools', '$0', 'GitHub Student Pack: Free access to JetBrains, GitHub Pro, etc.'),
        ('Testing & QA Tools', '$0', 'Free tools: Jest, Pytest, GitHub Actions for CI/CD'),
        ('Cloud Storage (Files/Images)', '$0', 'GitHub Student Pack: Free storage credits via AWS Educate or Azure'),
        ('Security Tools', '$0', 'GitHub Advanced Security, CodeQL (free for public/education repos)'),
        ('Project Management', '$0', 'GitHub Projects, Trello (free tier)'),
        ('Email Service', '$0', 'SendGrid (free tier: 100 emails/day)'),
        ('Monitoring & Analytics', '$0', 'Free tools: Google Analytics, Sentry (free tier)'),
        ('Documentation', '$0', 'GitHub Pages, GitBook (free tier)'),
        ('Contingency (Miscellaneous)', '$50', 'Buffer for unforeseen expenses'),
    ]
    
    # Add budget items as formatted text
    for item, cost, notes in budget_items:
        p = doc.add_paragraph()
        p.add_run(f'{item}: ').bold = True
        p.add_run(f'{cost}')
        if notes:
            p.add_run(f'\n   → {notes}').italic = True
    
    p = doc.add_paragraph()
    p.add_run('\nTotal First Year Budget: ~$62 USD').bold = True
    
    p = doc.add_paragraph()
    p.add_run(
        '\nGitHub Student Developer Pack Benefits:\n'
        'The project significantly benefits from the GitHub Student Developer Pack, which provides:\n'
    )
    doc.add_paragraph('$200 in DigitalOcean credits (2 years)', style='List Bullet')
    doc.add_paragraph('$100 in Microsoft Azure credits', style='List Bullet')
    doc.add_paragraph('Free domain name via Namecheap (1 year)', style='List Bullet')
    doc.add_paragraph('Free JetBrains Professional Developer Tools', style='List Bullet')
    doc.add_paragraph('GitHub Pro features (unlimited private repositories)', style='List Bullet')
    doc.add_paragraph('Free SSL certificates and CDN via Cloudflare', style='List Bullet')
    doc.add_paragraph('Access to premium development and deployment tools', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run(
        '\nNote: Using student benefits, the project can be developed and deployed with minimal '
        'financial investment, making it highly cost-effective and accessible for educational purposes.'
    )
    
    # 2.6 Feasibility Study
    doc.add_heading('2.6 Feasibility Study', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        'A comprehensive feasibility analysis has been conducted to determine the viability of '
        'the Rental Management System across technical, economic, operational, and legal dimensions.\n'
    )
    
    doc.add_heading('2.6.1 Technical Feasibility', level=3)
    p = doc.add_paragraph()
    p.add_run('Assessment: ').bold = True
    p.add_run('Highly Feasible\n\n')
    
    technical_factors = [
        'Technology Stack Availability: All required technologies (React, Django, MySQL) are open-source, well-documented, and widely supported',
        'Development Expertise: Team has access to extensive learning resources, tutorials, and community support for the chosen tech stack',
        'Infrastructure: GitHub Student Pack provides necessary hosting, database, and deployment infrastructure at no cost',
        'Scalability: Django and React are proven technologies capable of handling the expected user load and data volume',
        'Integration: REST API architecture allows easy integration with third-party services and future enhancements',
        'Development Tools: Free access to professional IDEs (JetBrains), version control (GitHub Pro), and CI/CD tools',
    ]
    
    for factor in technical_factors:
        doc.add_paragraph(factor, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('\nResult: ').bold = True
    p.add_run('The project is technically feasible with available resources and expertise.')
    
    doc.add_heading('2.6.2 Economic Feasibility', level=3)
    p = doc.add_paragraph()
    p.add_run('Assessment: ').bold = True
    p.add_run('Highly Feasible\n\n')
    
    economic_factors = [
        'Low Initial Investment: Total first-year cost of only $62 USD due to GitHub Student benefits',
        'Zero Hosting Costs: Free hosting credits eliminate major recurring expenses',
        'No License Fees: All software and tools are either free or included in student pack',
        'Scalable Costs: Expenses grow gradually with usage, allowing budget adjustment over time',
        'Revenue Potential: System can be monetized through subscription plans or license sales',
        'Educational Value: Serves as learning project with practical application, providing ROI in skills development',
        'Long-term Sustainability: Minimal maintenance costs with cloud-based infrastructure',
    ]
    
    for factor in economic_factors:
        doc.add_paragraph(factor, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('\nResult: ').bold = True
    p.add_run('The project is economically feasible with excellent cost-benefit ratio.')
    
    doc.add_heading('2.6.3 Operational Feasibility', level=3)
    p = doc.add_paragraph()
    p.add_run('Assessment: ').bold = True
    p.add_run('Feasible\n\n')
    
    operational_factors = [
        'User Acceptance: Property management is a common need with clear user benefits',
        'Market Demand: Growing digitization in property management creates market opportunity',
        'User Training: Intuitive interface design minimizes training requirements',
        'Support Team: Development team can provide initial support and maintenance',
        'Deployment: Cloud-based deployment allows easy access without installation',
        'Maintenance: Modular architecture facilitates updates and bug fixes',
        'Documentation: Comprehensive documentation supports user adoption and maintenance',
    ]
    
    for factor in operational_factors:
        doc.add_paragraph(factor, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('\nResult: ').bold = True
    p.add_run('The project is operationally feasible with proper planning and execution.')
    
    doc.add_heading('2.6.4 Legal and Risk Feasibility', level=3)
    p = doc.add_paragraph()
    p.add_run('Assessment: ').bold = True
    p.add_run('Feasible\n\n')
    
    legal_factors = [
        'Open Source Compliance: All technologies used have permissive licenses (MIT, BSD, Apache)',
        'Data Protection: System implements security best practices for data protection',
        'Privacy Compliance: User data handling follows GDPR principles and best practices',
        'Authentication Security: JWT-based authentication with secure password hashing',
        'Risk Mitigation: Regular backups, error handling, and monitoring reduce operational risks',
        'Terms of Service: Clear user agreements and privacy policies can be implemented',
        'Liability Management: System logs and audit trails provide accountability',
    ]
    
    for factor in legal_factors:
        doc.add_paragraph(factor, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('\nResult: ').bold = True
    p.add_run('The project is legally feasible with proper security and compliance measures.')
    
    doc.add_heading('2.6.5 Overall Feasibility Conclusion', level=3)
    p = doc.add_paragraph()
    p.add_run(
        'Based on comprehensive analysis across all dimensions, the Rental Management System is '
        'highly feasible and viable. The project benefits significantly from:\n'
    )
    
    doc.add_paragraph('GitHub Student Developer Pack providing essential infrastructure at no cost', style='List Bullet')
    doc.add_paragraph('Proven, well-supported technology stack', style='List Bullet')
    doc.add_paragraph('Clear market need and user benefits', style='List Bullet')
    doc.add_paragraph('Low financial risk with high educational value', style='List Bullet')
    doc.add_paragraph('Strong foundation for future growth and monetization', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run(
        '\nThe combination of minimal costs, available expertise, and proven technologies makes '
        'this an ideal project for both educational purposes and potential commercial deployment.'
    )
    
    doc.add_page_break()

def add_admin_use_cases(doc):
    """Add Admin use cases"""
    doc.add_heading('2.2.1 Admin Use Cases', level=3)
    
    # Use Case: Manage Users
    doc.add_heading('Use Case: Manage Users', level=4)
    
    doc.add_paragraph('Diagram:')
    uml = '''@startuml
Admin --> (Manage Users)
(Manage Users) ..> (Create User) : <<includes>>
(Manage Users) ..> (Update User) : <<includes>>
(Manage Users) ..> (Delete User) : <<includes>>
(Manage Users) ..> (View User List) : <<includes>>
@enduml'''
    doc.add_paragraph(uml, style='Intense Quote')
    
    p = doc.add_paragraph()
    p.add_run('Brief Description\n').bold = True
    p.add_run('The Admin creates, updates, deletes, and manages user accounts in the system.')
    
    p = doc.add_paragraph()
    p.add_run('Initial Step-By-Step Description\n').bold = True
    p.add_run('Before this use case can be initiated, the Admin has already logged into the system.\n')
    
    steps = [
        'Admin navigates to the User Management page',
        'System displays list of existing users',
        'Admin selects to create new user or update existing user',
        'System presents user form with fields (username, email, role, password)',
        'Admin fills in the information and submits',
        'System validates the information',
        'System creates/updates the user account',
        'System returns confirmation and updated user list'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    p = doc.add_paragraph()
    p.add_run('User Story\n').bold = True
    p.add_run('As an Admin, I want to manage user accounts so that I can control system access and assign appropriate roles.')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria\n').bold = True
    criteria = [
        'Admin can view a list of all users with their roles',
        'Admin can create new users with username, email, and role',
        'Admin can update existing user information',
        'Admin can delete users (with confirmation)',
        'System validates email format and username uniqueness',
        'Passwords are securely hashed before storage',
        'Success/error messages are displayed appropriately'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet 2')
    
    # Use Case: Monitor System
    doc.add_heading('Use Case: Monitor System', level=4)
    
    p = doc.add_paragraph()
    p.add_run('Brief Description\n').bold = True
    p.add_run('The Admin monitors system health, views logs, and checks system statistics.')
    
    p = doc.add_paragraph()
    p.add_run('User Story\n').bold = True
    p.add_run('As an Admin, I want to monitor system performance so that I can ensure optimal operation and quickly identify issues.')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria\n').bold = True
    criteria = [
        'Admin can view system dashboard with key metrics',
        'Admin can access system logs and error reports',
        'Admin can see total counts of properties, tenants, and active leases',
        'Admin can view recent system activities',
        'System displays response time and server status'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet 2')

def add_landlord_use_cases(doc):
    """Add Landlord use cases"""
    doc.add_heading('2.2.2 Landlord Use Cases', level=3)
    
    # Use Case: Manage Properties
    doc.add_heading('Use Case: Manage Properties', level=4)
    
    doc.add_paragraph('Diagram:')
    uml = '''@startuml
Landlord --> (Manage Properties)
(Manage Properties) ..> (Add Property) : <<includes>>
(Manage Properties) ..> (Update Property) : <<includes>>
(Manage Properties) ..> (Add Unit) : <<includes>>
(Manage Properties) ..> (Upload Images) : <<includes>>
@enduml'''
    doc.add_paragraph(uml, style='Intense Quote')
    
    p = doc.add_paragraph()
    p.add_run('Brief Description\n').bold = True
    p.add_run('The Landlord creates and manages rental properties and their associated units.')
    
    p = doc.add_paragraph()
    p.add_run('Initial Step-By-Step Description\n').bold = True
    
    steps = [
        'Landlord navigates to Properties page',
        'System displays list of landlord\'s properties',
        'Landlord selects to add new property',
        'System presents property form (name, address, type, description)',
        'Landlord fills in property details and uploads images',
        'System validates and saves property information',
        'Landlord adds units to the property',
        'System creates units with specified details (number, size, rent)',
        'System returns to property list with new property displayed'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    p = doc.add_paragraph()
    p.add_run('User Story\n').bold = True
    p.add_run('As a Landlord, I want to manage my properties and units so that I can effectively organize and track my rental inventory.')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria\n').bold = True
    criteria = [
        'Landlord can add new properties with complete details',
        'Landlord can upload multiple images per property',
        'Landlord can add multiple units to each property',
        'Each unit has unique identifier, size, and rent amount',
        'Landlord can edit property and unit information',
        'Landlord can mark properties/units as available or occupied',
        'System validates required fields before saving'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet 2')
    
    # Use Case: Manage Leases
    doc.add_heading('Use Case: Manage Leases', level=4)
    
    p = doc.add_paragraph()
    p.add_run('Brief Description\n').bold = True
    p.add_run('The Landlord creates and manages lease agreements with tenants.')
    
    p = doc.add_paragraph()
    p.add_run('User Story\n').bold = True
    p.add_run('As a Landlord, I want to create and manage lease agreements so that I can formalize rental arrangements with tenants.')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria\n').bold = True
    criteria = [
        'Landlord can create new lease with tenant and unit selection',
        'Lease includes start date, end date, and rent amount',
        'Landlord can upload lease documents (PDF/images)',
        'Landlord can view all active and expired leases',
        'Landlord can update lease information',
        'Landlord can terminate leases with reason',
        'System validates date ranges and tenant availability'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet 2')
    
    # Use Case: Track Payments
    doc.add_heading('Use Case: Track Payments', level=4)
    
    p = doc.add_paragraph()
    p.add_run('User Story\n').bold = True
    p.add_run('As a Landlord, I want to track rental payments so that I can monitor income and identify overdue accounts.')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria\n').bold = True
    criteria = [
        'Landlord can view all payments for their properties',
        'Landlord can see payment status (paid, pending, overdue)',
        'Landlord can record manual payments',
        'System automatically calculates total rent collected',
        'Landlord can filter payments by property, tenant, or date range',
        'System highlights overdue payments'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet 2')
    
    # Use Case: Manage Expenses
    doc.add_heading('Use Case: Manage Expenses', level=4)
    
    p = doc.add_paragraph()
    p.add_run('User Story\n').bold = True
    p.add_run('As a Landlord, I want to record property expenses so that I can track costs and calculate net income.')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria\n').bold = True
    criteria = [
        'Landlord can add expenses with description, amount, date, and category',
        'Landlord can attach receipts to expense records',
        'Landlord can view expense history',
        'System categorizes expenses (maintenance, utilities, taxes, etc.)',
        'Landlord can generate expense reports by property or time period'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet 2')

def add_tenant_use_cases(doc):
    """Add Tenant use cases"""
    doc.add_heading('2.2.3 Tenant Use Cases', level=3)
    
    # Use Case: View Lease Details
    doc.add_heading('Use Case: View Lease Details', level=4)
    
    doc.add_paragraph('Diagram:')
    uml = '''@startuml
Tenant --> (View Lease Details)
(View Lease Details) ..> (View Terms) : <<includes>>
(View Lease Details) ..> (Download Documents) : <<includes>>
@enduml'''
    doc.add_paragraph(uml, style='Intense Quote')
    
    p = doc.add_paragraph()
    p.add_run('Brief Description\n').bold = True
    p.add_run('The Tenant views their lease agreement details and downloads associated documents.')
    
    p = doc.add_paragraph()
    p.add_run('User Story\n').bold = True
    p.add_run('As a Tenant, I want to view my lease details so that I can understand my rental obligations and download my lease documents.')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria\n').bold = True
    criteria = [
        'Tenant can view complete lease information',
        'Tenant can see unit details and property information',
        'Tenant can download lease documents',
        'Tenant can view lease start and end dates',
        'Tenant can see monthly rent amount',
        'System displays lease status (active, expired, terminated)'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet 2')
    
    # Use Case: Make Payment
    doc.add_heading('Use Case: Make Payment', level=4)
    
    p = doc.add_paragraph()
    p.add_run('User Story\n').bold = True
    p.add_run('As a Tenant, I want to record my rent payments so that my payment history is tracked in the system.')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria\n').bold = True
    criteria = [
        'Tenant can submit payment information',
        'Tenant can specify payment amount and date',
        'Tenant can upload payment receipt',
        'Tenant can view payment history',
        'System calculates outstanding balance',
        'Tenant receives confirmation after payment submission'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet 2')
    
    # Use Case: Submit Notification
    doc.add_heading('Use Case: Submit Notification', level=4)
    
    p = doc.add_paragraph()
    p.add_run('User Story\n').bold = True
    p.add_run('As a Tenant, I want to send notifications to my landlord so that I can report issues or communicate important information.')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria\n').bold = True
    criteria = [
        'Tenant can create new notification with subject and message',
        'Tenant can categorize notification (maintenance, complaint, inquiry)',
        'Tenant can attach files to notifications',
        'Tenant can view sent notifications and their status',
        'System sends notification to property landlord',
        'Tenant receives responses from landlord'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet 2')

def add_requirements_specification(doc):
    """Add Section 3.0 - Requirements Specification"""
    doc.add_heading('3.0 Requirements Specification', level=1)
    
    # 3.1 External Interface Requirements
    doc.add_heading('3.1 External Interface Requirements', level=2)
    
    doc.add_heading('3.1.1 User Interfaces', level=3)
    p = doc.add_paragraph()
    p.add_run(
        'The system provides a responsive web interface built with React and Tailwind CSS. '
        'Key interface components include:\n'
    )
    interfaces = [
        'Login/Registration pages with form validation',
        'Role-based dashboards for Admin, Landlord, and Tenant',
        'Property listing and detail views with image galleries',
        'Data tables with sorting, filtering, and pagination',
        'Forms for creating/editing entities (properties, leases, etc.)',
        'Notification center for messages and alerts',
        'File upload interfaces for documents and images'
    ]
    for interface in interfaces:
        doc.add_paragraph(interface, style='List Bullet')
    
    doc.add_heading('3.1.2 Hardware Interfaces', level=3)
    p = doc.add_paragraph()
    p.add_run(
        'The system is web-based and requires:\n'
    )
    doc.add_paragraph('Client: Modern web browser on desktop or mobile device', style='List Bullet')
    doc.add_paragraph('Server: Linux/Windows server with Python 3.10+ support', style='List Bullet')
    doc.add_paragraph('Database: MySQL 8.0 server', style='List Bullet')
    
    doc.add_heading('3.1.3 Software Interfaces', level=3)
    p = doc.add_paragraph()
    p.add_run('The system interfaces with:\n')
    interfaces = [
        'Django 5.x - Web framework and ORM',
        'Django REST Framework - API development',
        'MySQL 8.0 - Relational database management',
        'JWT - Authentication token system',
        'React 18 - Frontend framework',
        'Axios - HTTP client for API calls'
    ]
    for interface in interfaces:
        doc.add_paragraph(interface, style='List Bullet')
    
    doc.add_heading('3.1.4 Communication Interfaces', level=3)
    p = doc.add_paragraph()
    p.add_run(
        'HTTP/HTTPS protocols for client-server communication. '
        'RESTful API with JSON data format. JWT tokens in Authorization headers '
        'for authenticated requests.'
    )
    
    # 3.2 Functional Requirements
    doc.add_heading('3.2 Functional Requirements', level=2)
    
    add_functional_requirement(doc, '3.2.1', 'User Authentication',
        'The system shall authenticate users using JWT tokens',
        ['User provides username and password',
         'System validates credentials against database',
         'System generates JWT token on successful authentication',
         'Token is stored in client browser',
         'Token is included in subsequent API requests'])
    
    add_functional_requirement(doc, '3.2.2', 'Property Management',
        'The system shall allow landlords to create and manage properties',
        ['Landlord provides property details (name, address, type)',
         'System validates required fields',
         'System stores property in database',
         'System associates property with landlord',
         'System allows image uploads for property'])
    
    add_functional_requirement(doc, '3.2.3', 'Unit Management',
        'The system shall allow management of units within properties',
        ['Landlord selects property to add unit',
         'System presents unit form',
         'Landlord provides unit details (number, size, rent)',
         'System validates and saves unit',
         'System links unit to parent property'])
    
    add_functional_requirement(doc, '3.2.4', 'Lease Creation',
        'The system shall enable creation of lease agreements',
        ['Landlord selects tenant and unit',
         'System presents lease form',
         'Landlord enters lease terms (dates, amount)',
         'Landlord uploads lease documents',
         'System creates lease record',
         'System marks unit as occupied'])
    
    add_functional_requirement(doc, '3.2.5', 'Payment Recording',
        'The system shall track rent and other payments',
        ['User creates payment entry',
         'System validates payment details',
         'System records payment with timestamp',
         'System updates tenant balance',
         'System generates payment confirmation'])
    
    add_functional_requirement(doc, '3.2.6', 'Expense Tracking',
        'The system shall record property-related expenses',
        ['Landlord enters expense details',
         'System validates expense information',
         'Landlord uploads receipt if available',
         'System categorizes expense',
         'System stores expense record'])
    
    add_functional_requirement(doc, '3.2.7', 'Document Management',
        'The system shall store and manage documents',
        ['User uploads document file',
         'System validates file type and size',
         'System stores file securely',
         'System creates document record with metadata',
         'User can download document later'])
    
    add_functional_requirement(doc, '3.2.8', 'Notification System',
        'The system shall provide communication between users',
        ['Tenant creates notification',
         'System validates notification content',
         'System sends to appropriate landlord',
         'Landlord receives and views notification',
         'Landlord can respond to notification'])
    
    # 3.3 Detailed Non-Functional Requirements
    doc.add_heading('3.3 Detailed Non-Functional Requirements', level=2)
    
    doc.add_heading('3.3.1 Performance Requirements', level=3)
    requirements = [
        'System shall respond to 95% of user requests within 2 seconds',
        'System shall support up to 10,000 concurrent users',
        'Database queries shall execute in less than 500ms',
        'Page load time shall not exceed 3 seconds on standard broadband',
        'API endpoints shall handle 1000 requests per second'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('3.3.2 Security Requirements', level=3)
    requirements = [
        'All passwords shall be hashed using industry-standard algorithms',
        'User sessions shall expire after 24 hours of inactivity',
        'API endpoints shall validate JWT tokens for authentication',
        'Role-based access control shall prevent unauthorized access',
        'All data transmission shall use HTTPS encryption',
        'System shall log all authentication attempts',
        'File uploads shall be validated for type and scanned for malware'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('3.3.3 Reliability Requirements', level=3)
    requirements = [
        'System shall have 99.5% uptime',
        'Database shall be backed up daily',
        'System shall recover from crashes within 5 minutes',
        'Data integrity checks shall run automatically',
        'Failed transactions shall be rolled back completely'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('3.3.4 Usability Requirements', level=3)
    requirements = [
        'Interface shall be accessible on desktop and mobile devices',
        'System shall provide clear error messages',
        'Forms shall include inline validation',
        'System shall follow WCAG 2.1 accessibility guidelines',
        'Common tasks shall require no more than 3 clicks',
        'Help documentation shall be available in-app'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('3.3.5 Maintainability Requirements', level=3)
    requirements = [
        'Code shall follow PEP 8 (Python) and ESLint (JavaScript) standards',
        'All functions shall include docstring documentation',
        'System shall have automated test coverage above 80%',
        'Code shall be modular and follow SOLID principles',
        'Database schema changes shall use migration scripts'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    # Add Data Model Section
    doc.add_page_break()
    doc.add_heading('3.4 Data Model', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        'The following diagram shows the logical structure of the Rental Management System database:'
    )
    
    doc.add_heading('Figure 6 - Data Model Structure (Class Diagram)', level=3)
    doc.add_paragraph('PlantUML Code:')
    
    uml_code = '''@startuml
!define ENTITY class

ENTITY User {
  +id: Integer
  +username: String
  +email: String
  +role: String
  +password: Hash
  +phone: String
  +profile_picture: File
  +created_at: DateTime
}

ENTITY Property {
  +id: Integer
  +name: String
  +address: String
  +type: String
  +description: Text
  +status: String
  +landlord_id: FK
  +created_at: DateTime
}

ENTITY PropertyImage {
  +id: Integer
  +property_id: FK
  +image: File
  +caption: String
}

ENTITY Unit {
  +id: Integer
  +property_id: FK
  +unit_number: String
  +bedrooms: Integer
  +bathrooms: Decimal
  +size: Decimal
  +rent_amount: Decimal
  +status: String
}

ENTITY Tenant {
  +id: Integer
  +user_id: FK
  +date_of_birth: Date
  +occupation: String
  +emergency_contact: String
  +emergency_phone: String
}

ENTITY Lease {
  +id: Integer
  +unit_id: FK
  +tenant_id: FK
  +start_date: Date
  +end_date: Date
  +rent_amount: Decimal
  +security_deposit: Decimal
  +status: String
  +created_at: DateTime
}

ENTITY Payment {
  +id: Integer
  +lease_id: FK
  +amount: Decimal
  +payment_date: Date
  +payment_method: String
  +status: String
  +receipt: File
}

ENTITY Expense {
  +id: Integer
  +property_id: FK
  +description: Text
  +amount: Decimal
  +expense_date: Date
  +category: String
  +receipt: File
}

ENTITY Document {
  +id: Integer
  +lease_id: FK
  +title: String
  +file: File
  +uploaded_at: DateTime
}

ENTITY Notification {
  +id: Integer
  +tenant_id: FK
  +landlord_id: FK
  +subject: String
  +message: Text
  +status: String
  +created_at: DateTime
}

' Relationships
User "1" -- "0..*" Property : owns >
Property "1" -- "0..*" PropertyImage : has >
Property "1" -- "0..*" Unit : contains >
Property "1" -- "0..*" Expense : incurs >
User "1" -- "0..1" Tenant : is >
Unit "1" -- "0..*" Lease : leased_by >
Tenant "1" -- "0..*" Lease : signs >
Lease "1" -- "0..*" Payment : receives >
Lease "1" -- "0..*" Document : has >
Tenant "1" -- "0..*" Notification : sends >
User "1" -- "0..*" Notification : receives >

@enduml'''
    doc.add_paragraph(uml_code, style='Intense Quote')
    
    # Add State Diagram
    doc.add_page_break()
    doc.add_heading('Figure 7 - Lease Management Process (State Diagram)', level=3)
    doc.add_paragraph('PlantUML Code:')
    
    state_uml = '''@startuml
[*] --> Draft
Draft --> Active : approve & sign
Draft --> Cancelled : cancel

Active --> Renewed : renew
Active --> Terminated : terminate
Active --> Expired : end_date reached

Renewed --> Active : process renewal
Expired --> [*]
Terminated --> [*]
Cancelled --> [*]

note right of Draft
  Lease created but not yet active
end note

note right of Active
  Tenant occupying unit
  Payments being tracked
end note

note right of Expired
  Lease term completed
end note

note right of Terminated
  Early termination by either party
end note

@enduml'''
    doc.add_paragraph(state_uml, style='Intense Quote')
    
    doc.add_page_break()

def add_functional_requirement(doc, number, title, description, steps):
    """Helper function to add a functional requirement"""
    doc.add_heading(f'{number} {title}', level=3)
    p = doc.add_paragraph()
    p.add_run('Description: ').bold = True
    p.add_run(description)
    
    p = doc.add_paragraph()
    p.add_run('Processing Steps:\n').bold = True
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')

def add_appendix(doc):
    """Add appendix section"""
    doc.add_heading('Appendix A: PlantUML Diagram Instructions', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        'To render the PlantUML diagrams in this document:\n\n'
        '1. Visit PlantUML Online Editor: http://www.plantuml.com/plantuml/\n'
        '2. Copy the PlantUML code from any diagram section\n'
        '3. Paste the code into the editor\n'
        '4. The diagram will be rendered automatically\n'
        '5. You can download the diagram as PNG, SVG, or other formats\n\n'
        'Alternatively, install PlantUML locally:\n'
        '- Install Java (required for PlantUML)\n'
        '- Download PlantUML JAR file\n'
        '- Run: java -jar plantuml.jar diagram.puml\n'
    )

def create_srs_document():
    """Main function to create the SRS document"""
    
    # Create document
    doc = Document()
    
    # Set up document properties
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title page
    title = doc.add_heading('Software Requirements Specification', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_heading('Rental Management System', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    
    # Team information
    team_info = [
        'Development Team',
        'Full-Stack Web Application',
        '',
        'Submitted for:',
        'Software Engineering Project',
        '',
        f'Date: {datetime.now().strftime("%B %d, %Y")}',
        '',
        'Version 1.0'
    ]
    
    for line in team_info:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Add all sections
    add_toc(doc)
    add_figures_list(doc)
    add_introduction(doc)
    add_overall_description(doc)
    add_requirements_specification(doc)
    add_appendix(doc)
    
    # Save document
    output_path = '/home/runner/work/software_project/software_project/SRS_RentalManagementSystem.docx'
    doc.save(output_path)
    print(f'SRS Document created successfully: {output_path}')
    return output_path

if __name__ == '__main__':
    create_srs_document()
