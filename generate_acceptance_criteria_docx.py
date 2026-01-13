#!/usr/bin/env python3
"""
Script to convert ACCEPTANCE_CRITERIA.md to DOCX format.
Creates a professionally formatted Microsoft Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re

def add_horizontal_line(paragraph):
    """Add a horizontal line to a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_acceptance_criteria_docx():
    """Generate DOCX document from acceptance criteria"""
    
    # Create document
    doc = Document()
    
    # Set up document properties
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title page
    title = doc.add_heading('Acceptance Criteria', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Rental Management System', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    # Document info
    info_text = [
        'Software Requirements Specification',
        'Acceptance Criteria Document',
        '',
        f'Version: 1.0',
        f'Date: {datetime.now().strftime("%B %d, %Y")}',
        'Status: Final'
    ]
    
    for line in info_text:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Admin Features',
        '2. Landlord Features',
        '3. Tenant Features',
        '4. System-Wide Features',
        '5. Document Control'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    p = doc.add_paragraph()
    p.add_run(
        'This document defines the acceptance criteria for each feature in the Rental Management System. '
        'These criteria are used to determine when a feature is complete and ready for production. '
        'Each criterion is testable, traceable, and linked to specific user stories.'
    )
    
    doc.add_paragraph()
    
    # Admin Features
    add_admin_criteria(doc)
    
    # Landlord Features
    add_landlord_criteria(doc)
    
    # Tenant Features
    add_tenant_criteria(doc)
    
    # System Features
    add_system_criteria(doc)
    
    # Document Control
    add_document_control(doc)
    
    # Save document
    output_path = '/home/runner/work/software_project/software_project/ACCEPTANCE_CRITERIA.docx'
    doc.save(output_path)
    print(f'Acceptance Criteria DOCX created: {output_path}')
    return output_path

def add_admin_criteria(doc):
    """Add admin acceptance criteria"""
    doc.add_page_break()
    doc.add_heading('1. Admin Features', level=1)
    
    # AC-A001: Manage Users
    doc.add_heading('AC-A001: Manage Users', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('User account management for system administrators')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    criteria = [
        ('User List Display', [
            'Admin can view a list of all users in the system',
            'List displays username, email, role, and status for each user',
            'List supports sorting by username, email, role, or creation date',
            'List includes pagination with configurable items per page'
        ]),
        ('Create New User', [
            'Admin can create new users with username, email, password, and role',
            'System validates email format (must be valid email address)',
            'System validates username uniqueness (no duplicates allowed)',
            'Password must meet minimum requirements (8+ characters, mixed case, numbers)',
            'Role must be one of: Admin, Landlord, or Tenant',
            'Success message is displayed upon successful creation'
        ]),
        ('Update Existing User', [
            'Admin can update user information (email, role, status)',
            'Admin can reset user password',
            'System prevents admin from removing their own admin role',
            'Changes are logged in the system audit trail',
            'Success message is displayed upon successful update'
        ]),
        ('Delete User', [
            'Admin can delete user accounts',
            'System displays confirmation dialog before deletion',
            'System prevents deletion of currently logged-in admin',
            'System shows warning if user has associated data (properties, leases)',
            'Deleted users are soft-deleted (marked as inactive, not removed from database)'
        ]),
        ('Search and Filter', [
            'Admin can search users by username or email',
            'Admin can filter users by role',
            'Admin can filter users by active/inactive status',
            'Search results update in real-time as user types'
        ]),
        ('Password Security', [
            'All passwords are hashed using bcrypt or equivalent',
            'Passwords are never stored in plain text',
            'Passwords are never displayed in the UI',
            'Password reset generates secure token sent via email'
        ]),
        ('Error Handling', [
            'System displays clear error messages for validation failures',
            'System displays error if email is already in use',
            'System displays error if username is already taken',
            'Network errors are handled gracefully with retry option'
        ])
    ]
    
    for i, (title, items) in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')
    
    # AC-A002: Monitor System
    doc.add_heading('AC-A002: Monitor System', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('System health monitoring and statistics dashboard')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    criteria = [
        ('Dashboard Display', [
            'Admin can view system dashboard with key metrics',
            'Dashboard displays total number of properties',
            'Dashboard displays total number of active leases',
            'Dashboard displays total number of tenants',
            'Dashboard displays total number of landlords',
            'Dashboard refreshes automatically every 5 minutes'
        ]),
        ('System Logs', [
            'Admin can access system logs',
            'Logs include timestamp, user, action, and result',
            'Logs support filtering by date range',
            'Logs support filtering by user',
            'Logs support filtering by action type',
            'Admin can export logs to CSV format'
        ]),
        ('Performance Metrics', [
            'Dashboard displays current server response time',
            'Dashboard displays database query performance',
            'Dashboard displays memory usage',
            'Dashboard displays active user sessions',
            'Performance metrics are updated every minute'
        ]),
        ('Error Reports', [
            'Admin can view recent system errors',
            'Error reports include stack trace when available',
            'Error reports show affected user and timestamp',
            'Admin can mark errors as resolved',
            'Critical errors trigger email notification to admins'
        ]),
        ('Activity Timeline', [
            'Dashboard shows recent system activities',
            'Timeline displays user logins and logouts',
            'Timeline displays data modifications',
            'Timeline displays file uploads',
            'Timeline is limited to last 100 activities'
        ])
    ]
    
    for i, (title, items) in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')

def add_landlord_criteria(doc):
    """Add landlord acceptance criteria"""
    doc.add_page_break()
    doc.add_heading('2. Landlord Features', level=1)
    
    # AC-L001: Manage Properties
    doc.add_heading('AC-L001: Manage Properties', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('Create and manage rental properties with units and images')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    criteria = [
        ('Property Creation', [
            'Landlord can add new properties with name, address, type, and description',
            'Property type can be: Residential, Commercial, or Mixed Use',
            'Address must include street, city, state, and zip code',
            'System validates required fields before saving',
            'Success message is displayed after property creation',
            'Newly created property appears in landlord\'s property list'
        ]),
        ('Property Image Upload', [
            'Landlord can upload multiple images per property (minimum 1, maximum 10)',
            'Supported image formats: JPG, PNG, WebP',
            'Maximum file size: 5MB per image',
            'Images are automatically resized to standard dimensions',
            'Landlord can set one image as the primary/featured image',
            'Landlord can delete uploaded images'
        ]),
        ('Add Units to Property', [
            'Landlord can add multiple units to each property',
            'Each unit requires: unit number, bedrooms, bathrooms, size (sq ft), rent amount',
            'Unit number must be unique within the property',
            'Rent amount must be a positive decimal number',
            'Bedrooms and bathrooms can be integers or decimals (e.g., 2.5 bathrooms)',
            'Unit status can be: Available, Occupied, or Under Maintenance'
        ]),
        ('Update Property Information', [
            'Landlord can edit property details at any time',
            'Changes to address are validated',
            'Changes to property name are allowed',
            'System preserves property creation timestamp',
            'Update history is maintained in audit log'
        ]),
        ('Update Unit Information', [
            'Landlord can edit unit details',
            'Landlord can change unit status',
            'Landlord can update rent amount',
            'Rent updates do not affect existing leases',
            'Changes are effective immediately for new leases'
        ]),
        ('Property Status Management', [
            'Landlord can mark property as Available or Unavailable',
            'Unavailable properties are hidden from tenant searches',
            'Landlord can view count of available vs occupied units',
            'System prevents marking property unavailable if active leases exist'
        ]),
        ('Validation and Error Handling', [
            'System validates all required fields',
            'System prevents duplicate property names for same landlord',
            'System displays clear error messages for validation failures',
            'System validates numeric fields (rent, size) are positive numbers',
            'System handles image upload failures gracefully'
        ])
    ]
    
    for i, (title, items) in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')
    
    # Add other landlord criteria (L002, L003, L004)
    add_landlord_lease_criteria(doc)
    add_landlord_payment_criteria(doc)
    add_landlord_expense_criteria(doc)

def add_landlord_lease_criteria(doc):
    """Add landlord lease management criteria"""
    doc.add_heading('AC-L002: Manage Leases', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('Create and manage lease agreements with tenants')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    criteria = [
        ('Lease Creation', [
            'Landlord can create new lease agreement',
            'Landlord selects tenant from list of registered tenants',
            'Landlord selects available unit from their properties',
            'Landlord enters start date and end date',
            'Landlord enters monthly rent amount',
            'Landlord enters security deposit amount',
            'Start date must be today or in the future',
            'End date must be after start date'
        ]),
        ('Lease Document Upload', [
            'Landlord can upload lease agreement document (PDF format)',
            'Maximum file size: 10MB',
            'Multiple documents can be attached to one lease',
            'Document titles are required for each upload',
            'Documents can be downloaded later by both parties'
        ]),
        ('Lease Status Tracking', [
            'System automatically sets lease status to "Active" on start date',
            'System automatically sets lease status to "Expired" on end date',
            'Landlord can manually terminate lease before end date',
            'Landlord must provide reason for early termination',
            'Termination requires confirmation dialog'
        ]),
        ('Lease Renewal', [
            'Landlord can renew existing lease',
            'Renewal creates new lease record linked to original',
            'Landlord can modify terms during renewal',
            'New rent amount can be different from original',
            'Renewal updates unit status appropriately'
        ]),
        ('View Lease History', [
            'Landlord can view all leases for their properties',
            'Landlord can filter by property',
            'Landlord can filter by status (Active, Expired, Terminated)',
            'Landlord can filter by tenant name',
            'Lease list shows key details: tenant, unit, dates, status'
        ]),
        ('Lease Validation', [
            'System validates tenant availability (not already in active lease)',
            'System validates unit availability (not already leased)',
            'System validates date ranges are logical',
            'System validates rent and deposit are positive numbers',
            'System prevents overlapping leases for same unit'
        ]),
        ('Unit Status Synchronization', [
            'When lease becomes active, unit status changes to "Occupied"',
            'When lease ends/terminates, unit status changes to "Available"',
            'Status changes are automatic and real-time',
            'Landlord receives notification of status changes'
        ])
    ]
    
    for i, (title, items) in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')

def add_landlord_payment_criteria(doc):
    """Add landlord payment tracking criteria"""
    doc.add_heading('AC-L003: Track Payments', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('Monitor and record rental payments from tenants')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    criteria = [
        ('Payment List View', [
            'Landlord can view all payments for their properties',
            'Payment list shows: tenant name, property, unit, amount, date, status',
            'Payments are sorted by date (newest first) by default',
            'List includes pagination for large datasets',
            'Landlord can export payment list to CSV'
        ]),
        ('Payment Status Indicators', [
            'Each payment shows status: Paid, Pending, or Overdue',
            'Overdue payments are highlighted in red',
            'Pending payments are highlighted in yellow',
            'Paid payments show green checkmark',
            'Status is calculated automatically based on due date'
        ]),
        ('Manual Payment Recording', [
            'Landlord can manually record payments received',
            'Landlord enters payment amount and payment date',
            'Landlord selects payment method (Cash, Check, Bank Transfer, Other)',
            'Landlord can attach receipt image',
            'System validates amount is positive number',
            'System links payment to correct lease'
        ]),
        ('Payment Filtering', [
            'Landlord can filter payments by property',
            'Landlord can filter by tenant name',
            'Landlord can filter by date range',
            'Landlord can filter by payment status',
            'Landlord can filter by payment method',
            'Multiple filters can be applied simultaneously'
        ]),
        ('Payment Calculations', [
            'System automatically calculates total rent collected per property',
            'System calculates total rent collected per month',
            'System calculates outstanding balance per tenant',
            'System calculates collection rate percentage',
            'Calculations update in real-time when payments are added'
        ]),
        ('Overdue Payment Identification', [
            'System identifies payments not received by due date',
            'Overdue payments are flagged in payment list',
            'System calculates days overdue for each late payment',
            'Landlord can view total overdue amount',
            'Landlord can view list of tenants with overdue payments'
        ]),
        ('Payment Notifications', [
            'System sends payment reminder to tenant 3 days before due date',
            'System notifies landlord when payment is submitted by tenant',
            'System notifies landlord of overdue payments',
            'Notifications can be configured in settings'
        ])
    ]
    
    for i, (title, items) in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')

def add_landlord_expense_criteria(doc):
    """Add landlord expense management criteria"""
    doc.add_heading('AC-L004: Manage Expenses', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('Record and track property-related expenses')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    criteria = [
        ('Expense Entry', [
            'Landlord can add new expense records',
            'Landlord enters description, amount, date, and category',
            'Categories include: Maintenance, Utilities, Taxes, Insurance, Repairs, Other',
            'Landlord selects associated property',
            'Amount must be positive decimal number',
            'Date can be past or current date (not future)'
        ]),
        ('Receipt Upload', [
            'Landlord can attach receipt image or PDF to expense',
            'Supported formats: JPG, PNG, PDF',
            'Maximum file size: 5MB',
            'Receipt attachment is optional',
            'Multiple receipts can be attached to one expense',
            'Receipts can be viewed and downloaded later'
        ]),
        ('Expense Viewing', [
            'Landlord can view list of all expenses',
            'Expense list shows: date, description, amount, category, property',
            'List is sorted by date (newest first) by default',
            'List includes pagination',
            'Landlord can view expense details including receipt'
        ]),
        ('Expense Filtering', [
            'Landlord can filter expenses by property',
            'Landlord can filter by category',
            'Landlord can filter by date range',
            'Landlord can filter by amount range',
            'Multiple filters work together'
        ]),
        ('Expense Reporting', [
            'Landlord can generate expense report by date range',
            'Report shows total expenses by category',
            'Report shows total expenses by property',
            'Report can be exported to PDF or CSV',
            'Report includes summary statistics'
        ]),
        ('Expense Categories', [
            'System provides predefined expense categories',
            'Landlord can select "Other" and provide custom description',
            'Category totals are calculated automatically',
            'Categories help with tax reporting',
            'Most-used categories are shown first in dropdown'
        ]),
        ('Edit and Delete Expenses', [
            'Landlord can edit expense details',
            'Landlord can delete expenses with confirmation',
            'Edit history is maintained in audit log',
            'Deleted expenses are soft-deleted (archived, not removed)',
            'Original expense data is preserved for reporting'
        ])
    ]
    
    for i, (title, items) in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')

def add_tenant_criteria(doc):
    """Add tenant acceptance criteria"""
    doc.add_page_break()
    doc.add_heading('3. Tenant Features', level=1)
    
    # AC-T001: View Lease Details
    doc.add_heading('AC-T001: View Lease Details', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('Access and review lease agreement information')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    criteria = [
        ('Lease Information Display', [
            'Tenant can view complete lease details',
            'Display shows property name and address',
            'Display shows unit number and details',
            'Display shows lease start and end dates',
            'Display shows monthly rent amount',
            'Display shows security deposit amount',
            'Display shows lease status (Active, Expired, Terminated)'
        ]),
        ('Property and Unit Information', [
            'Tenant can view property description',
            'Tenant can view property images',
            'Tenant can view unit size (square footage)',
            'Tenant can view number of bedrooms and bathrooms',
            'Tenant can view landlord contact information',
            'Information is read-only (tenant cannot edit)'
        ]),
        ('Document Access', [
            'Tenant can view list of lease documents',
            'Tenant can download lease agreement PDF',
            'Tenant can download other associated documents',
            'Document titles are displayed clearly',
            'Download button is prominently displayed'
        ]),
        ('Lease Timeline', [
            'Tenant can see visual timeline of lease period',
            'Timeline shows start date, current date, and end date',
            'System calculates and displays days remaining in lease',
            'System calculates and displays percentage of lease completed',
            'Timeline updates automatically each day'
        ]),
        ('Lease Status', [
            'Current status is displayed prominently',
            'Status includes clear visual indicator (icon/color)',
            'Active leases show green indicator',
            'Expired leases show gray indicator',
            'Terminated leases show red indicator with reason'
        ]),
        ('Landlord Information', [
            'Tenant can view landlord name',
            'Tenant can view landlord email',
            'Tenant can view landlord phone number',
            'Contact information is clickable (email opens mail client)',
            'Information is always current and accurate'
        ])
    ]
    
    for i, (title, items) in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')
    
    # Add other tenant criteria
    add_tenant_payment_criteria(doc)
    add_tenant_notification_criteria(doc)

def add_tenant_payment_criteria(doc):
    """Add tenant payment criteria"""
    doc.add_heading('AC-T002: Make Payment', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('Submit and track rental payment information')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    criteria = [
        ('Payment Submission', [
            'Tenant can submit payment information',
            'Tenant enters payment amount',
            'Tenant enters payment date',
            'Tenant selects payment method from dropdown',
            'Payment methods: Bank Transfer, Cash, Check, Online Payment',
            'System validates amount is positive number',
            'System validates date is not in future'
        ]),
        ('Receipt Upload', [
            'Tenant can upload payment receipt (optional)',
            'Supported formats: JPG, PNG, PDF',
            'Maximum file size: 5MB',
            'System displays upload progress',
            'System confirms successful upload',
            'Tenant can view uploaded receipt after submission'
        ]),
        ('Payment Confirmation', [
            'System displays confirmation message after submission',
            'Confirmation includes payment ID for reference',
            'Confirmation includes submitted amount and date',
            'Tenant receives email confirmation',
            'Payment appears in payment history immediately'
        ]),
        ('Payment History View', [
            'Tenant can view all their payment history',
            'History shows: date, amount, method, status',
            'History is sorted by date (newest first)',
            'Each payment shows status indicator',
            'Tenant can filter by date range'
        ]),
        ('Outstanding Balance', [
            'System calculates and displays current balance owed',
            'Balance calculation includes rent, late fees, other charges',
            'Balance is updated in real-time when payments are submitted',
            'Tenant can see next payment due date',
            'Tenant can see next payment amount due'
        ]),
        ('Payment Status Tracking', [
            'Each payment shows current status',
            'Status options: Submitted, Pending Review, Confirmed, Rejected',
            'Status is updated by landlord/admin',
            'Tenant receives notification when status changes',
            'Status history is maintained'
        ]),
        ('Validation and Error Handling', [
            'System validates required fields before submission',
            'System prevents duplicate submissions (double-click protection)',
            'Error messages are clear and actionable',
            'Failed submissions can be retried',
            'Network errors are handled gracefully'
        ])
    ]
    
    for i, (title, items) in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')

def add_tenant_notification_criteria(doc):
    """Add tenant notification criteria"""
    doc.add_heading('AC-T003: Submit Notification', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('Send messages and notifications to landlord')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    criteria = [
        ('Notification Creation', [
            'Tenant can create new notification/message',
            'Tenant enters subject line (required, max 100 characters)',
            'Tenant enters message body (required, max 2000 characters)',
            'Tenant selects notification type/category',
            'Categories: Maintenance Request, Complaint, Inquiry, General',
            'Subject and message use text input with character counter'
        ]),
        ('File Attachment', [
            'Tenant can attach files to notification (optional)',
            'Supported formats: JPG, PNG, PDF, DOC, DOCX',
            'Maximum 3 files per notification',
            'Maximum 5MB per file',
            'System displays upload progress for each file',
            'Tenant can remove attached files before sending'
        ]),
        ('Notification Submission', [
            'Tenant clicks Send button to submit',
            'System validates all required fields',
            'System displays confirmation dialog',
            'System shows success message after submission',
            'Notification appears in sent items immediately',
            'Landlord receives notification in their inbox'
        ]),
        ('Sent Notifications View', [
            'Tenant can view all sent notifications',
            'List shows: date, subject, category, status',
            'Status options: Sent, Read, Responded, Resolved',
            'List is sorted by date (newest first)',
            'Tenant can click to view notification details'
        ]),
        ('Notification Status Tracking', [
            'Each notification shows current status with icon',
            '"Sent" - delivered to landlord',
            '"Read" - opened by landlord',
            '"Responded" - landlord replied',
            '"Resolved" - issue marked as complete',
            'Status updates automatically without page refresh'
        ]),
        ('Response Viewing', [
            'Tenant can view landlord responses',
            'Responses appear as threaded conversation',
            'Each response shows timestamp and sender',
            'Tenant receives email notification of new responses',
            'Response history is maintained'
        ]),
        ('Validation and Limits', [
            'Subject line is required (cannot be empty)',
            'Message body is required (minimum 10 characters)',
            'Category selection is required',
            'System prevents submission of empty notifications',
            'Clear error messages guide tenant to fix issues',
            'Rate limiting prevents spam (max 10 notifications per day)'
        ])
    ]
    
    for i, (title, items) in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')

def add_system_criteria(doc):
    """Add system-wide acceptance criteria"""
    doc.add_page_break()
    doc.add_heading('4. System-Wide Features', level=1)
    
    # AC-S001: User Authentication
    doc.add_heading('AC-S001: User Authentication', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('Secure login and session management')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    criteria = [
        ('User Login', [
            'User can log in with username and password',
            'System validates credentials against database',
            'System generates JWT token on successful login',
            'Token expires after 24 hours',
            'Failed login attempts are logged',
            'Account locks after 5 failed attempts in 15 minutes'
        ]),
        ('User Registration', [
            'New users can register with email, username, password',
            'System validates email format',
            'System checks username uniqueness',
            'Password requirements: minimum 8 characters, mixed case, number',
            'Email verification sent to new users',
            'Account activated after email verification'
        ]),
        ('Password Reset', [
            'Users can request password reset via email',
            'System generates secure reset token',
            'Reset link expires after 1 hour',
            'User can set new password via reset link',
            'Old password is invalidated after reset',
            'User receives confirmation email after reset'
        ]),
        ('Session Management', [
            'JWT token stored securely in browser',
            'Token included in all API requests',
            'Session expires after 24 hours of inactivity',
            'User redirected to login on session expiry',
            'Logout clears token from browser'
        ]),
        ('Role-Based Access Control', [
            'User role determines accessible features',
            'Admin can access all features',
            'Landlord can access property and tenant management',
            'Tenant can access their lease and payment info',
            'Unauthorized access attempts are blocked with 403 error'
        ])
    ]
    
    for i, (title, items) in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')
    
    # Add remaining system criteria
    doc.add_heading('AC-S002: Data Validation', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('Input validation and data integrity')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    validation_criteria = [
        ('Client-Side Validation', [
            'Required fields are validated before form submission',
            'Email format validated using regex',
            'Numeric fields validated for correct data type',
            'Date fields validated for logical ranges',
            'Error messages displayed inline with fields'
        ]),
        ('Server-Side Validation', [
            'All inputs re-validated on server',
            'SQL injection attempts are blocked',
            'XSS attempts are sanitized',
            'File uploads validated for type and size',
            'Malicious payloads rejected with 400 error'
        ]),
        ('Data Constraints', [
            'Unique constraints enforced on database level',
            'Foreign key relationships maintained',
            'Numeric values constrained to valid ranges',
            'Dates validated for logical order',
            'String lengths enforced at database level'
        ])
    ]
    
    for i, (title, items) in enumerate(validation_criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')
    
    # Performance criteria
    doc.add_heading('AC-S003: Performance', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('System responsiveness and scalability')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    perf_criteria = [
        ('Response Time', [
            '95% of API requests respond within 2 seconds',
            'Page load time under 3 seconds on standard broadband',
            'Database queries execute in under 500ms',
            'Large lists implement pagination',
            'Images are lazy-loaded for better performance'
        ]),
        ('Scalability', [
            'System supports 10,000 concurrent users',
            'Database indexed on frequently queried fields',
            'API endpoints can handle 1000 requests/second',
            'Static assets cached by browser',
            'Database connection pooling implemented'
        ])
    ]
    
    for i, (title, items) in enumerate(perf_criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')
    
    # Security criteria
    doc.add_heading('AC-S004: Security', level=2)
    p = doc.add_paragraph()
    p.add_run('Feature: ').bold = True
    p.add_run('Data protection and secure communication')
    
    doc.add_heading('Acceptance Criteria:', level=3)
    
    sec_criteria = [
        ('Data Encryption', [
            'All passwords hashed with bcrypt',
            'HTTPS used for all communications',
            'Sensitive data encrypted at rest',
            'JWT tokens signed with secret key',
            'Database credentials stored securely in environment variables'
        ]),
        ('Access Control', [
            'Role-based permissions enforced on all endpoints',
            'Users can only access their own data',
            'Admins have full system access',
            'Failed authorization logged',
            'SQL injection prevented by parameterized queries'
        ]),
        ('Audit Trail', [
            'All data modifications logged',
            'Logs include user, timestamp, action',
            'Logs retained for 90 days',
            'Admin can review audit logs',
            'Critical actions require confirmation'
        ])
    ]
    
    for i, (title, items) in enumerate(sec_criteria, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet 2')

def add_document_control(doc):
    """Add document control section"""
    doc.add_page_break()
    doc.add_heading('5. Document Control', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Version: ').bold = True
    p.add_run('1.0\n')
    p.add_run('Date: ').bold = True
    p.add_run(f'{datetime.now().strftime("%B %d, %Y")}\n')
    p.add_run('Status: ').bold = True
    p.add_run('Final\n')
    p.add_run('Author: ').bold = True
    p.add_run('Development Team\n')
    
    doc.add_heading('Revision History', level=2)
    
    # Add table for revision history
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    headers = ['Version', 'Date', 'Author', 'Description']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data row
    row = table.rows[1]
    row.cells[0].text = '1.0'
    row.cells[1].text = datetime.now().strftime('%b %d, %Y')
    row.cells[2].text = 'Dev Team'
    row.cells[3].text = 'Initial acceptance criteria document'
    
    doc.add_paragraph()
    
    doc.add_heading('Usage Notes', level=2)
    
    p = doc.add_paragraph()
    p.add_run('For Developers:\n').bold = True
    notes = [
        'Use these criteria to understand when a feature is complete',
        'Implement unit tests that verify each criterion',
        'Reference AC codes in pull requests (e.g., "Implements AC-L001.1-7")'
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('For Testers:\n').bold = True
    notes = [
        'Create test cases based on these criteria',
        'Each criterion should have at least one test case',
        'Mark criteria as pass/fail during testing',
        'Report any failed criteria as bugs'
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('For Product Owners:\n').bold = True
    notes = [
        'Use these criteria during sprint planning',
        'Reference during feature demos and acceptance',
        'Update criteria as requirements evolve',
        'Ensure criteria align with business goals'
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')

if __name__ == '__main__':
    create_acceptance_criteria_docx()
