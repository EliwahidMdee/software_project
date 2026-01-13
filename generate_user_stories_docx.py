#!/usr/bin/env python3
"""
Script to convert USER_STORIES.md to DOCX format.
Creates a professionally formatted Microsoft Word document with embedded PlantUML code.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_user_stories_docx():
    """Generate DOCX document from user stories"""
    
    # Create document
    doc = Document()
    
    # Set up document properties
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title page
    title = doc.add_heading('User Stories', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Rental Management System', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    # Document info
    info_text = [
        'Software Requirements Specification',
        'User Stories with Use Case Diagrams',
        '',
        f'Version: 1.0',
        f'Date: {datetime.now().strftime("%B %d, %Y")}',
        'Status: Final'
    ]
    
    for line in info_text:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Admin User Stories',
        '2. Landlord User Stories',
        '3. Tenant User Stories',
        '4. System User Stories',
        '5. Story Point Reference',
        '6. Document Control'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    p = doc.add_paragraph()
    p.add_run(
        'This document contains detailed user stories for each feature in the Rental Management System, '
        'organized by user role. Each section includes PlantUML use case diagrams that can be rendered '
        'at http://www.plantuml.com/plantuml/'
    )
    
    doc.add_paragraph()
    
    # Add sections
    add_admin_stories(doc)
    add_landlord_stories(doc)
    add_tenant_stories(doc)
    add_system_stories(doc)
    add_reference_section(doc)
    add_document_control_section(doc)
    
    # Save document
    output_path = '/home/runner/work/software_project/software_project/USER_STORIES.docx'
    doc.save(output_path)
    print(f'User Stories DOCX created: {output_path}')
    return output_path

def add_admin_stories(doc):
    """Add admin user stories"""
    doc.add_page_break()
    doc.add_heading('1. Admin User Stories', level=1)
    
    # Add use case diagram
    doc.add_heading('Use Case Diagram: Admin Role', level=2)
    
    uml_code = '''@startuml
' Admin Use Cases - Complete View

left to right direction

actor "Admin" as Admin

package "Rental Management System" {
    rectangle "User Management" {
        usecase "Manage Users" as UC_A1
        usecase "Create User" as UC_A1a
        usecase "Update User" as UC_A1b
        usecase "Delete User" as UC_A1c
        usecase "View User List" as UC_A1d
        usecase "Reset Password" as UC_A1e
        usecase "Assign Roles" as UC_A1f
    }
    
    rectangle "System Monitoring" {
        usecase "Monitor System" as UC_A2
        usecase "View Dashboard" as UC_A2a
        usecase "Access Logs" as UC_A2b
        usecase "View Statistics" as UC_A2c
        usecase "Check Performance" as UC_A2d
        usecase "View Error Reports" as UC_A2e
    }
}

Admin --> UC_A1
Admin --> UC_A2

UC_A1 ..> UC_A1a : <<includes>>
UC_A1 ..> UC_A1b : <<includes>>
UC_A1 ..> UC_A1c : <<includes>>
UC_A1 ..> UC_A1d : <<includes>>
UC_A1 ..> UC_A1e : <<includes>>
UC_A1 ..> UC_A1f : <<includes>>

UC_A2 ..> UC_A2a : <<includes>>
UC_A2 ..> UC_A2b : <<includes>>
UC_A2 ..> UC_A2c : <<includes>>
UC_A2 ..> UC_A2d : <<includes>>
UC_A2 ..> UC_A2e : <<includes>>

@enduml'''
    
    doc.add_paragraph('PlantUML Code:')
    doc.add_paragraph(uml_code, style='Intense Quote')
    
    # User Story A-001
    doc.add_heading('User Story A-001: Create User Account', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As an ').italic = True
    p.add_run('Admin').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('create new user accounts in the system').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('authorized users can access the rental management platform').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High\n')
    p.add_run('Story Points: ').bold = True
    p.add_run('3\n')
    p.add_run('Sprint: ').bold = True
    p.add_run('1\n')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-A001.2 in ACCEPTANCE_CRITERIA document')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tasks:').bold = True
    tasks = [
        'Design user creation form UI',
        'Implement form validation',
        'Create API endpoint for user creation',
        'Implement password hashing',
        'Add email uniqueness check',
        'Write unit tests',
        'Write integration tests'
    ]
    for task in tasks:
        doc.add_paragraph(f'☐ {task}', style='List Bullet')
    
    # Add more admin stories
    add_admin_story_002(doc)
    add_admin_story_003(doc)
    add_admin_story_004(doc)
    add_admin_story_005(doc)

def add_admin_story_002(doc):
    """Add Admin Story 002"""
    doc.add_heading('User Story A-002: Update User Information', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As an ').italic = True
    p.add_run('Admin').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('update existing user account information').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can keep user records accurate and current').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('3  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('1')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-A001.3 in ACCEPTANCE_CRITERIA document')
    
    tasks = [
        'Design user edit form UI',
        'Implement form pre-population with current data',
        'Create API endpoint for user updates',
        'Implement audit logging',
        'Prevent self-demotion for admins',
        'Write unit tests'
    ]
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tasks:').bold = True
    for task in tasks:
        doc.add_paragraph(f'☐ {task}', style='List Bullet')

def add_admin_story_003(doc):
    """Add Admin Story 003"""
    doc.add_heading('User Story A-003: Delete User Account', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As an ').italic = True
    p.add_run('Admin').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('delete user accounts').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can remove inactive or unauthorized users from the system').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('Medium  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('2  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('2')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-A001.4 in ACCEPTANCE_CRITERIA document')

def add_admin_story_004(doc):
    """Add Admin Story 004"""
    doc.add_heading('User Story A-004: View System Dashboard', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As an ').italic = True
    p.add_run('Admin').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('view a system dashboard with key metrics').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can monitor system health and usage at a glance').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('Medium  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('5  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('2')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-A002.1 in ACCEPTANCE_CRITERIA document')

def add_admin_story_005(doc):
    """Add Admin Story 005"""
    doc.add_heading('User Story A-005: Access System Logs', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As an ').italic = True
    p.add_run('Admin').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('view and filter system logs').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can troubleshoot issues and audit system activity').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('Medium  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('5  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('3')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-A002.2 in ACCEPTANCE_CRITERIA document')

def add_landlord_stories(doc):
    """Add landlord user stories"""
    doc.add_page_break()
    doc.add_heading('2. Landlord User Stories', level=1)
    
    # Add use case diagram
    doc.add_heading('Use Case Diagram: Landlord Role', level=2)
    
    uml_code = '''@startuml
' Landlord Use Cases - Complete View

left to right direction

actor "Landlord" as Landlord

package "Rental Management System" {
    rectangle "Property Management" {
        usecase "Manage Properties" as UC_L1
        usecase "Add Property" as UC_L1a
        usecase "Add Unit" as UC_L1d
        usecase "Upload Images" as UC_L1f
    }
    
    rectangle "Lease Management" {
        usecase "Manage Leases" as UC_L2
        usecase "Create Lease" as UC_L2a
        usecase "Terminate Lease" as UC_L2c
        usecase "Upload Documents" as UC_L2e
    }
    
    rectangle "Payment Tracking" {
        usecase "Track Payments" as UC_L3
        usecase "View Payments" as UC_L3a
        usecase "Record Payment" as UC_L3b
    }
    
    rectangle "Expense Management" {
        usecase "Manage Expenses" as UC_L4
        usecase "Add Expense" as UC_L4a
        usecase "Upload Receipt" as UC_L4d
    }
}

Landlord --> UC_L1
Landlord --> UC_L2
Landlord --> UC_L3
Landlord --> UC_L4

UC_L1 ..> UC_L1a : <<includes>>
UC_L1 ..> UC_L1d : <<includes>>
UC_L1 ..> UC_L1f : <<includes>>

UC_L2 ..> UC_L2a : <<includes>>
UC_L2 ..> UC_L2c : <<includes>>
UC_L2 ..> UC_L2e : <<includes>>

UC_L3 ..> UC_L3a : <<includes>>
UC_L3 ..> UC_L3b : <<includes>>

UC_L4 ..> UC_L4a : <<includes>>
UC_L4 ..> UC_L4d : <<includes>>

@enduml'''
    
    doc.add_paragraph('PlantUML Code:')
    doc.add_paragraph(uml_code, style='Intense Quote')
    
    # Add landlord stories (abbreviated for space)
    add_landlord_story_001(doc)
    add_landlord_story_002(doc)
    add_landlord_story_003(doc)
    add_landlord_story_004(doc)

def add_landlord_story_001(doc):
    """Add Landlord Story 001"""
    doc.add_heading('User Story L-001: Add New Property', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('Landlord').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('add new rental properties to the system').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can manage them and make them available for tenants').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('5  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('1')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-L001.1 in ACCEPTANCE_CRITERIA document')

def add_landlord_story_002(doc):
    """Add Landlord Story 002"""
    doc.add_heading('User Story L-002: Upload Property Images', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('Landlord').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('upload multiple images for each property').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('potential tenants can see what the property looks like').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('5  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('1')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-L001.2 in ACCEPTANCE_CRITERIA document')

def add_landlord_story_003(doc):
    """Add Landlord Story 003"""
    doc.add_heading('User Story L-003: Add Units to Property', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('Landlord').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('add multiple rental units to each property').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can track individual apartments or rooms separately').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('5  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('1')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-L001.3 in ACCEPTANCE_CRITERIA document')

def add_landlord_story_004(doc):
    """Add Landlord Story 004"""
    doc.add_heading('User Story L-004: Create Lease Agreement', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('Landlord').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('create lease agreements with tenants').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can formalize rental arrangements and track occupancy').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('8  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('2')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-L002.1 in ACCEPTANCE_CRITERIA document')

def add_tenant_stories(doc):
    """Add tenant user stories"""
    doc.add_page_break()
    doc.add_heading('3. Tenant User Stories', level=1)
    
    # Add use case diagram
    doc.add_heading('Use Case Diagram: Tenant Role', level=2)
    
    uml_code = '''@startuml
' Tenant Use Cases - Complete View

left to right direction

actor "Tenant" as Tenant

package "Rental Management System" {
    rectangle "Lease Information" {
        usecase "View Lease Details" as UC_T1
        usecase "View Terms" as UC_T1a
        usecase "Download Documents" as UC_T1d
    }
    
    rectangle "Payment Management" {
        usecase "Make Payment" as UC_T2
        usecase "Submit Payment Info" as UC_T2a
        usecase "Upload Receipt" as UC_T2b
        usecase "View Payment History" as UC_T2c
    }
    
    rectangle "Communication" {
        usecase "Submit Notification" as UC_T3
        usecase "Create Message" as UC_T3a
        usecase "Attach Files" as UC_T3b
        usecase "View Responses" as UC_T3d
    }
}

Tenant --> UC_T1
Tenant --> UC_T2
Tenant --> UC_T3

UC_T1 ..> UC_T1a : <<includes>>
UC_T1 ..> UC_T1d : <<includes>>

UC_T2 ..> UC_T2a : <<includes>>
UC_T2 ..> UC_T2b : <<includes>>
UC_T2 ..> UC_T2c : <<includes>>

UC_T3 ..> UC_T3a : <<includes>>
UC_T3 ..> UC_T3b : <<includes>>
UC_T3 ..> UC_T3d : <<includes>>

@enduml'''
    
    doc.add_paragraph('PlantUML Code:')
    doc.add_paragraph(uml_code, style='Intense Quote')
    
    # Add tenant stories
    add_tenant_story_001(doc)
    add_tenant_story_002(doc)
    add_tenant_story_003(doc)

def add_tenant_story_001(doc):
    """Add Tenant Story 001"""
    doc.add_heading('User Story T-001: View My Lease Details', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('Tenant').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('view my current lease agreement details').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can understand my rental obligations and rights').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('3  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('2')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-T001.1 in ACCEPTANCE_CRITERIA document')

def add_tenant_story_002(doc):
    """Add Tenant Story 002"""
    doc.add_heading('User Story T-002: Download Lease Documents', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('Tenant').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('download my lease agreement and related documents').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I have a copy for my records').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('2  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('2')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-T001.3 in ACCEPTANCE_CRITERIA document')

def add_tenant_story_003(doc):
    """Add Tenant Story 003"""
    doc.add_heading('User Story T-003: Submit Rent Payment', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('Tenant').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('submit my rent payment information').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('my landlord knows I\'ve paid and can verify the payment').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('5  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('2')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-T002.1 in ACCEPTANCE_CRITERIA document')

def add_system_stories(doc):
    """Add system user stories"""
    doc.add_page_break()
    doc.add_heading('4. System User Stories', level=1)
    
    # Add use case diagram
    doc.add_heading('Use Case Diagram: System-Wide Features', level=2)
    
    uml_code = '''@startuml
' System-Wide Use Cases

left to right direction

actor "User" as User

package "Authentication & Security" {
    usecase "Register Account" as UC_S1
    usecase "Login" as UC_S2
    usecase "Logout" as UC_S3
    usecase "Reset Password" as UC_S4
    usecase "Change Password" as UC_S6
}

package "Profile Management" {
    usecase "View Profile" as UC_S7
    usecase "Update Profile" as UC_S8
    usecase "Upload Photo" as UC_S9
}

User --> UC_S1
User --> UC_S2
User --> UC_S3
User --> UC_S4
User --> UC_S6
User --> UC_S7
User --> UC_S8
User --> UC_S9

@enduml'''
    
    doc.add_paragraph('PlantUML Code:')
    doc.add_paragraph(uml_code, style='Intense Quote')
    
    # Add system stories
    add_system_story_001(doc)
    add_system_story_002(doc)
    add_system_story_003(doc)

def add_system_story_001(doc):
    """Add System Story 001"""
    doc.add_heading('User Story S-001: Register New Account', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('new user').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('register for an account').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can access the rental management system').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('5  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('1')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-S001.2 in ACCEPTANCE_CRITERIA document')

def add_system_story_002(doc):
    """Add System Story 002"""
    doc.add_heading('User Story S-002: Login to System', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('registered user').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('login with my credentials').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can access my account and features').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('3  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('1')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-S001.1 in ACCEPTANCE_CRITERIA document')

def add_system_story_003(doc):
    """Add System Story 003"""
    doc.add_heading('User Story S-003: Reset Forgotten Password', level=2)
    
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('user who forgot their password').bold = True
    p.add_run('\nI want to ').italic = True
    p.add_run('reset my password via email').bold = True
    p.add_run('\nSo that ').italic = True
    p.add_run('I can regain access to my account').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('High  |  ')
    p.add_run('Story Points: ').bold = True
    p.add_run('5  |  ')
    p.add_run('Sprint: ').bold = True
    p.add_run('1')
    
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria: ').bold = True
    p.add_run('See AC-S001.3 in ACCEPTANCE_CRITERIA document')

def add_reference_section(doc):
    """Add story point reference"""
    doc.add_page_break()
    doc.add_heading('5. Story Point Reference', level=1)
    
    doc.add_paragraph('Story points are used to estimate the complexity and effort required for each user story.')
    
    doc.add_paragraph()
    
    # Add table
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    table.rows[0].cells[0].text = 'Story Points'
    table.rows[0].cells[1].text = 'Effort Estimate'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Data rows
    data = [
        ('1 point', 'Very simple task, < 4 hours'),
        ('2 points', 'Simple task, 4-8 hours'),
        ('3 points', 'Moderate task, 1-2 days'),
        ('5 points', 'Complex task, 2-3 days'),
        ('8 points', 'Very complex task, 3-5 days'),
        ('13 points', 'Epic, should be broken down')
    ]
    
    for i, (points, effort) in enumerate(data, 1):
        table.rows[i].cells[0].text = points
        table.rows[i].cells[1].text = effort
    
    doc.add_paragraph()
    
    doc.add_heading('Priority Levels', level=2)
    
    priorities = [
        ('High', 'Critical for MVP, must be in first release'),
        ('Medium', 'Important but can be delayed if necessary'),
        ('Low', 'Nice to have, can be added in later versions')
    ]
    
    for priority, description in priorities:
        p = doc.add_paragraph()
        p.add_run(f'{priority}: ').bold = True
        p.add_run(description)

def add_document_control_section(doc):
    """Add document control"""
    doc.add_page_break()
    doc.add_heading('6. Document Control', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Version: ').bold = True
    p.add_run('1.0\n')
    p.add_run('Date: ').bold = True
    p.add_run(f'{datetime.now().strftime("%B %d, %Y")}\n')
    p.add_run('Status: ').bold = True
    p.add_run('Final\n')
    p.add_run('Author: ').bold = True
    p.add_run('Development Team\n')
    
    doc.add_heading('How to Use This Document', level=2)
    
    p = doc.add_paragraph()
    p.add_run('For Product Owners:\n').bold = True
    items = [
        'Use to plan sprints and prioritize features',
        'Reference during backlog grooming',
        'Use story points for velocity planning'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('For Developers:\n').bold = True
    items = [
        'Reference during implementation',
        'Use as basis for creating subtasks',
        'Link commits to story IDs'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('For QA:\n').bold = True
    items = [
        'Create test cases based on acceptance criteria',
        'Verify each story meets all criteria',
        'Reference during test planning'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Rendering PlantUML Diagrams', level=2)
    
    p = doc.add_paragraph()
    p.add_run('To view the use case diagrams:\n\n')
    steps = [
        'Copy the PlantUML code (including @startuml and @enduml)',
        'Visit http://www.plantuml.com/plantuml/',
        'Paste the code into the editor',
        'The diagram will render automatically',
        'Download as PNG or SVG if needed'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

if __name__ == '__main__':
    create_user_stories_docx()
